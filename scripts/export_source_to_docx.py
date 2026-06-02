from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT_DIR = REPO_ROOT / "docs" / "source_exports"

SKIP_DIRS = {
    ".git",
    ".github",
    ".agents",
    ".claude",
    ".ruff_cache",
    ".next",
    ".turbo",
    ".vite",
    "coverage",
    "dist",
    "build",
    "node_modules",
    "test-artifacts",
    "uploads",
    "__pycache__",
}

SKIP_FILENAMES = {
    "package-lock.json",
    "yarn.lock",
    "pnpm-lock.yaml",
    "bun.lockb",
    "skills-lock.json",
}

SKIP_SUFFIXES = {
    ".bmp",
    ".cache",
    ".class",
    ".dll",
    ".docx",
    ".exe",
    ".gif",
    ".ico",
    ".jpeg",
    ".jpg",
    ".lock",
    ".log",
    ".map",
    ".pdf",
    ".png",
    ".pyc",
    ".sqlite",
    ".sqlite3",
    ".webp",
    ".zip",
}

INCLUDE_SUFFIXES = {
    ".css",
    ".cjs",
    ".html",
    ".js",
    ".json",
    ".mjs",
    ".prisma",
    ".sql",
    ".toml",
    ".ts",
    ".tsx",
    ".yml",
    ".yaml",
}

SECRET_ASSIGNMENT_PATTERN = re.compile(
    r"""(?ix)
    ^
    (?P<prefix>\s*
      (?:
        (?:export\s+)?(?:const|let|var)\s+|
      )?
      (?P<key>[\w.-]*(?:secret|token|password|passwd|pwd|credential|api[_-]?key|private[_-]?key)[\w.-]*)
      \s*[:=]\s*
    )
    (?P<quote>['"])
    (?P<value>[^'"]{4,})
    (?P=quote)
    (?P<suffix>.*)$
    """
)

ENV_ACCESS_PATTERN = re.compile(
    r"""(?ix)
    (?P<prefix>process\.env\.[A-Z0-9_]*(?:SECRET|TOKEN|PASSWORD|PASSWD|PWD|CREDENTIAL|API_KEY|PRIVATE_KEY)[A-Z0-9_]*\s*
      (?:\|\||\?\?)\s*
    )
    (?P<quote>['"])
    (?P<value>[^'"]{4,})
    (?P=quote)
    """
)

URL_CREDENTIAL_PATTERN = re.compile(
    r"(?P<scheme>[a-z][a-z0-9+.-]*://)(?P<user>[^/\s:@]+):(?P<password>[^@\s/]+)@",
    re.IGNORECASE,
)

SECRETISH_VALUE_PATTERN = re.compile(
    r"""(?ix)
    (
      ^[a-z0-9_+./~=-]{16,}$
      | ^[a-z0-9_-]+\.[a-z0-9_-]+\.[a-z0-9_-]+$
      | ^sk-[a-z0-9_-]+$
      | ^[a-z0-9+/]{20,}={0,2}$
      | (?=\S+$)(?=.*\d)(?=.*[a-z])(?=.*[A-Z])(?=.*[^a-z0-9]).{10,}
    )
    """
)


@dataclass(frozen=True)
class ExportBucket:
    name: str
    output_name: str
    roots: tuple[Path, ...]
    extra_files: tuple[Path, ...] = ()


BUCKETS = (
    ExportBucket(
        name="Frontend",
        output_name="bankflow-frontend-source.docx",
        roots=(REPO_ROOT / "frontend",),
    ),
    ExportBucket(
        name="Backend",
        output_name="bankflow-backend-source.docx",
        roots=(REPO_ROOT / "backend" / "src",),
        extra_files=(
            REPO_ROOT / "backend" / "package.json",
            REPO_ROOT / "backend" / "tsconfig.json",
            REPO_ROOT / "backend" / "Dockerfile",
            REPO_ROOT / "backend" / "README.md",
        ),
    ),
    ExportBucket(
        name="Database and Prisma",
        output_name="bankflow-database-prisma-source.docx",
        roots=(REPO_ROOT / "backend" / "prisma", REPO_ROOT / "docker"),
        extra_files=(REPO_ROOT / "docker-compose.yml",),
    ),
)


def relative_path(path: Path) -> str:
    return path.relative_to(REPO_ROOT).as_posix()


def should_skip(path: Path) -> bool:
    rel_parts = path.relative_to(REPO_ROOT).parts
    if any(part in SKIP_DIRS for part in rel_parts):
        return True
    if path.name in SKIP_FILENAMES:
        return True
    if path.name.startswith(".env"):
        return True
    if path.suffix.lower() in SKIP_SUFFIXES:
        return True
    return path.suffix.lower() not in INCLUDE_SUFFIXES


def collect_files(bucket: ExportBucket) -> list[Path]:
    files: set[Path] = set()

    for root in bucket.roots:
        if not root.exists():
            continue
        if root.is_file():
            if not should_skip(root):
                files.add(root)
            continue
        for path in root.rglob("*"):
            if path.is_file() and not should_skip(path):
                files.add(path)

    for path in bucket.extra_files:
        if path.exists() and path.is_file() and not should_skip(path):
            files.add(path)

    return sorted(files, key=relative_path)


def redact_line(line: str) -> str:
    def redact_secret_assignment(match: re.Match[str]) -> str:
        key = match.group("key").lower()
        value = match.group("value")
        always_redact = any(word in key for word in ("password", "passwd", "pwd", "private_key"))
        should_redact = always_redact and " " not in value
        should_redact = should_redact or bool(SECRETISH_VALUE_PATTERN.search(value))
        if not should_redact:
            return match.group(0)
        return f"{match.group('prefix')}{match.group('quote')}[REDACTED]{match.group('quote')}{match.group('suffix')}"

    line = SECRET_ASSIGNMENT_PATTERN.sub(
        redact_secret_assignment,
        line,
    )
    line = ENV_ACCESS_PATTERN.sub(
        lambda match: f"{match.group('prefix')}{match.group('quote')}[REDACTED]{match.group('quote')}",
        line,
    )
    return URL_CREDENTIAL_PATTERN.sub(r"\g<scheme>\g<user>:[REDACTED]@", line)


def read_redacted(path: Path) -> str:
    text = path.read_text(encoding="utf-8", errors="replace")
    return "\n".join(redact_line(line.rstrip("\n")) for line in text.splitlines())


def add_code_block(document: Document, code: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

    run = paragraph.add_run(code if code else "[empty file]")
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(35, 39, 47)


def add_document_styles(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 2"].font.name = "Aptos"
    styles["Heading 2"].font.size = Pt(12)


def write_docx(bucket: ExportBucket, files: list[Path], output_dir: Path) -> Path:
    document = Document()
    add_document_styles(document)

    document.add_heading(f"Bankflow {bucket.name} Source Export", level=1)
    document.add_paragraph(
        "Generated from repository source files. Environment files, dependency folders, "
        "build artifacts, logs, binary assets, lockfiles, and likely credential literals are excluded or redacted."
    )
    document.add_paragraph(f"Files included: {len(files)}")

    document.add_heading("Included Files", level=2)
    for file_path in files:
        document.add_paragraph(relative_path(file_path), style="List Bullet")

    if files:
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    for index, file_path in enumerate(files):
        if index:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        document.add_heading(relative_path(file_path), level=2)
        add_code_block(document, read_redacted(file_path))

    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / bucket.output_name
    document.save(output_path)
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Export Bankflow source snippets into categorized DOCX files.")
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=DEFAULT_OUTPUT_DIR,
        help=f"Directory for generated DOCX files. Defaults to {DEFAULT_OUTPUT_DIR}",
    )
    args = parser.parse_args()

    output_dir = args.output_dir.resolve()
    for bucket in BUCKETS:
        files = collect_files(bucket)
        output_path = write_docx(bucket, files, output_dir)
        print(f"{bucket.name}: wrote {len(files)} files to {output_path}")


if __name__ == "__main__":
    main()
